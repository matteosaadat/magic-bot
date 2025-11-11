# AI INSTRUCTION:
# Recursively walk input folders, extract readable text from many file types,
# and write normalized .txt siblings into an "extracted" output root while
# mirroring the original folder structure.
#
# Supported:
# - .txt, .md, .markdown           (read as text)
# - .pdf                           (pdfminer.six)
# - .docx                          (python-docx)
# - .rtf                           (striprtf)
# - .html, .htm                    (beautifulsoup4)
# - .csv                           (csv module)
# - .xlsx                          (openpyxl)
# - Images: .png .jpg .jpeg .tiff  (optional OCR via pytesseract + pillow)
#
# Safety:
# - Skips files above max size
# - UTF-8 read with errors="ignore" fallbacks
# - Logs skipped/failed files
#
# Output:
#   For each input file, write a .txt file under OUT_ROOT mirroring paths.
#   Example:
#     src/data/raw/resume/Mahdi.pdf → src/data/tmp/extracted/resume/Mahdi.pdf.txt

from __future__ import annotations
import argparse
import csv
import io
import os
from pathlib import Path
from typing import Iterable, Optional

# Optional imports guarded inside functions
# pdfminer.six, python-docx, bs4, striprtf, openpyxl, pillow, pytesseract

TEXT_EXTS = {".txt", ".md", ".markdown"}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
RTF_EXTS = {".rtf"}
HTML_EXTS = {".html", ".htm"}
CSV_EXTS = {".csv"}
XLSX_EXTS = {".xlsx"}
IMG_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff"}

DEFAULT_ALLOWED = (
    TEXT_EXTS | PDF_EXTS | DOCX_EXTS | RTF_EXTS |
    HTML_EXTS | CSV_EXTS | XLSX_EXTS | IMG_EXTS
)

def _read_text_file(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="ignore")

def _extract_pdf(p: Path) -> str:
    from pdfminer.high_level import extract_text
    return extract_text(str(p))

def _extract_docx(p: Path) -> str:
    from docx import Document  # python-docx
    doc = Document(str(p))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # tables (lightweight)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)

def _extract_rtf(p: Path) -> str:
    from striprtf.striprtf import rtf_to_text
    data = p.read_text(encoding="utf-8", errors="ignore")
    return rtf_to_text(data)

def _extract_html(p: Path) -> str:
    from bs4 import BeautifulSoup
    html = p.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")
    # Remove script/style
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)

def _extract_csv(p: Path) -> str:
    with p.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        rows = []
        for row in reader:
            rows.append("\t".join(row))
        return "\n".join(rows)

def _extract_xlsx(p: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(filename=str(p), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [(str(v) if v is not None else "") for v in row]
            parts.append("\t".join(vals))
    return "\n".join(parts)

def _extract_image_ocr(p: Path, lang: str = "eng") -> Optional[str]:
    # Requires: Tesseract binary installed on system + pip: pillow, pytesseract
    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return None  # OCR not available; caller will log
    try:
        img = Image.open(str(p))
        return pytesseract.image_to_string(img, lang=lang)
    except Exception:
        return None

def _iter_files(inputs: Iterable[str], follow_symlinks: bool) -> Iterable[Path]:
    seen = set()
    for root in inputs:
        for p in Path(root).rglob("*"):
            if p.is_file():
                # dedupe by inode if possible
                key = (p.stat().st_ino, p.stat().st_dev) if hasattr(os, "stat") else (str(p),)
                if key in seen:
                    continue
                seen.add(key)
                # Follow symlinks if asked; otherwise skip them
                if p.is_symlink() and not follow_symlinks:
                    continue
                yield p

def extract_many(
    inputs: Iterable[str],
    out_root: str,
    allowed_exts = DEFAULT_ALLOWED,
    max_mb: float = 32.0,
    follow_symlinks: bool = False,
    ocr_images: bool = False,
    ocr_lang: str = "eng",
) -> dict:
    out_base = Path(out_root)
    out_base.mkdir(parents=True, exist_ok=True)

    stats = {
        "processed": 0,
        "written": 0,
        "skipped_size": 0,
        "skipped_ext": 0,
        "failed": 0,
        "ocr_used": 0,
    }

    for p in _iter_files(inputs, follow_symlinks=follow_symlinks):
        ext = p.suffix.lower()
        if ext not in allowed_exts:
            stats["skipped_ext"] += 1
            continue
        # size guard
        try:
            size_mb = p.stat().st_size / (1024 * 1024)
        except Exception:
            size_mb = 0
        if size_mb > max_mb:
            print(f"!! Skipping (size>{max_mb}MB): {p}")
            stats["skipped_size"] += 1
            continue

        stats["processed"] += 1

        rel = None
        for root in inputs:
            try:
                rel = p.relative_to(Path(root))
                break
            except ValueError:
                continue
        if rel is None:
            rel = p.name

        # Mirror folders; write .txt
        out_path = out_base / Path(rel)
        out_path = out_path.with_suffix(out_path.suffix + ".txt")
        out_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            text = ""
            if ext in TEXT_EXTS:
                text = _read_text_file(p)
            elif ext in PDF_EXTS:
                text = _extract_pdf(p)
            elif ext in DOCX_EXTS:
                text = _extract_docx(p)
            elif ext in RTF_EXTS:
                text = _extract_rtf(p)
            elif ext in HTML_EXTS:
                text = _extract_html(p)
            elif ext in CSV_EXTS:
                text = _extract_csv(p)
            elif ext in XLSX_EXTS:
                text = _extract_xlsx(p)
            elif ext in IMG_EXTS:
                if ocr_images:
                    ocr_text = _extract_image_ocr(p)
                    if ocr_text:
                        text = ocr_text
                        stats["ocr_used"] += 1
                    else:
                        print(f"!! OCR unavailable/failed: {p}")
                        text = ""  # still create empty placeholder to keep structure
                else:
                    print(f"!! Skipping image (OCR disabled): {p}")
                    text = ""
            else:
                # Shouldn't reach here due to allowed_exts filter
                text = ""

            # Write even if empty: keeps provenance and path mirror
            out_path.write_text(text or "", encoding="utf-8")
            stats["written"] += 1
            print(f">> {p}  →  {out_path}  ({len(text)} chars)")
        except Exception as e:
            stats["failed"] += 1
            print(f"!! Failed to extract {p}: {e}")

    return stats

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--inputs", nargs="+", required=True, help="One or more input folders")
    ap.add_argument("--out", required=True, help="Output root for extracted .txt")
    ap.add_argument("--max-mb", type=float, default=32.0)
    ap.add_argument("--follow-symlinks", action="store_true")
    ap.add_argument("--ocr-images", action="store_true", help="Enable OCR for images (requires Tesseract binary + pytesseract + pillow)")
    ap.add_argument("--ocr-lang", default="eng")
    args = ap.parse_args()

    stats = extract_many(
        inputs=args.inputs,
        out_root=args.out,
        max_mb=args.max_mb,
        follow_symlinks=args.follow_symlinks,
        ocr_images=args.ocr_images,
        ocr_lang=args.ocr_lang,
    )
    print("== Extraction stats ==", stats)

if __name__ == "__main__":
    main()
